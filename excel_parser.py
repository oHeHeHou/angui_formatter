import os

import xlrd
from docx import Document

content_header_index = 0
choice_header_index = 0
answer_header_index = 0

for root, dirs, files in os.walk('.'):
    for file in files:
        if file.endswith('xls'):
            print('文件名:{0}'.format(file))
            suffix_index = file.rindex('.xls')
            file_prefix = file[0:suffix_index]

            book = xlrd.open_workbook(file)
            print("Sheet数量： {0}".format(book.nsheets))
            print("Sheet名称: {0}".format(book.sheet_names()))

            document = Document()
            document.add_heading(file_prefix, 0)

            for i, sh in enumerate(book.sheets()):
                sh_name = sh.name
                # 写入word
                document.add_heading(sh_name, level=1)
                for rx in range(sh.nrows):
                    if rx == 0:
                        row_content = sh.row(rx)

                        # 记录表头中对应列的序号，用于后面获取内容
                        for index, row in enumerate(row_content):
                            header = row.value
                            if '题目' in header:
                                content_header_index = index
                            elif '选项' in header:
                                choice_header_index = index
                            elif '正确答案' in header:
                                answer_header_index = index
                        continue
                    row_content = sh.row(rx)
                    # 题目
                    row_subject = str(row_content[content_header_index].value).strip()
                    # 判断题后面加括号
                    if '判断' in sh_name:
                        try:
                            stop_index = row_subject.rindex('。')
                            row_subject = row_subject[0:stop_index]
                            row_subject += '（     ）'
                        except Exception:
                            row_subject += '（     ）'
                    else:
                        row_subject = row_subject.replace('（）', '(    )')
                        row_subject = row_subject.replace('()', '(    )')
                    row_choice = str(row_content[choice_header_index].value).strip()
                    # 选项
                    row_choice = row_choice.replace('|', '   ')
                    # 答案
                    row_answer = str(row_content[answer_header_index].value).strip()
                    if '判断' in sh_name:
                        if 'A' in row_answer:
                            row_answer = '正确'
                        elif 'B' in row_answer:
                            row_answer = '错误'
                    row_hint = '答案： ' + row_answer
                    # 写入题目
                    document.add_paragraph(row_subject, style='List Number')
                    if '判断' not in sh_name:
                        document.add_paragraph(row_choice)
                    # 写入答案
                    document.add_paragraph(row_hint)
                    out_file_name = file_prefix + '.docx'
                    document.save(out_file_name)
